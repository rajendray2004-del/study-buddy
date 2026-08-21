import streamlit as st
import pdfplumber
import google.generativeai as genai
import json
import datetime
st.set_page_config(page_title="Hey, Happy To See You ♥", page_icon="📚", layout="wide")
col1, col2 = st.columns([2, 1])
with col1:
    st.title("📚 AI Study Buddy")
    st.markdown("*Turn your notes into smart, adaptive quizzes*")
with col2:
    st.image("undraw_anime.svg", width=160)
st.markdown("""
    <style>
    @keyframes gradientShift {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }

    .stApp {
        background: linear-gradient(-45deg, #0f2027, #203a43, #2c5364, #0f2027);
        background-size: 400% 400%;
        animation: gradientShift 15s ease infinite;
    }

    h1 {
        background: linear-gradient(90deg, #00d2ff 0%, #3a7bd5 50%, #00d2ff 100%);
        background-size: 200% auto;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 800;
        animation: gradientShift 4s linear infinite;
    }

    h2, h3 {
        color: #64ffda;
    }

    p, label, .stMarkdown {
        color: #e0e0e0;
    }

    .stButton>button {
        background: linear-gradient(90deg, #00d2ff, #3a7bd5);
        color: white;
        border-radius: 30px;
        border: none;
        padding: 12px 28px;
        font-weight: 600;
        letter-spacing: 0.5px;
        transition: all 0.3s ease;
        box-shadow: 0 4px 10px rgba(0, 210, 255, 0.2);
    }

    .stButton>button:hover {
        transform: translateY(-3px) scale(1.03);
        box-shadow: 0 8px 20px rgba(0, 210, 255, 0.4);
    }

    div[data-testid="stRadio"] label {
        background-color: rgba(255,255,255,0.04);
        border: 1px solid rgba(100, 255, 218, 0.15);
        border-radius: 10px;
        padding: 10px 14px;
        margin: 5px 0;
        transition: all 0.25s ease;
    }

    div[data-testid="stRadio"] label:hover {
        background-color: rgba(100, 255, 218, 0.1);
        border-color: #64ffda;
        transform: translateX(4px);
    }

    .stFileUploader {
        border: 2px dashed #3a7bd5;
        border-radius: 14px;
        padding: 12px;
        background-color: rgba(255,255,255,0.02);
    }

    div[data-testid="stAlert"] {
        border-radius: 12px;
    }

    .stSlider {
        padding-top: 10px;
    }
        .stApp::before {
        content: "📚 ✏️ 🎓 📖 ✨ 📚 🎓 ✏️";
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        font-size: 40px;
        opacity: 0.06;
        line-height: 150px;
        letter-spacing: 60px;
        pointer-events: none;
        z-index: 0;
        animation: floatBg 20s linear infinite;
        overflow: hidden;
    }

    @keyframes floatBg {
        0% { transform: translateY(0px); }
        50% { transform: translateY(-30px); }
        100% { transform: translateY(0px); }
    }

    .main .block-container {
        position: relative;
        z-index: 1;
    }
    </style>
""", unsafe_allow_html=True)
st.markdown("""
    <style>
    .mascot {
        position: fixed;
        bottom: 20px;
        right: 20px;
        width: 90px;
        z-index: 999;
        animation: bounce 3s ease-in-out infinite;
    }
    @keyframes bounce {
        0%, 100% { transform: translateY(0px); }
        50% { transform: translateY(-15px); }
    }
    </style>

    <svg class="mascot" viewBox="0 0 200 200" xmlns="http://www.w3.org/2000/svg">
        <ellipse cx="100" cy="120" rx="65" ry="70" fill="#3a7bd5"/>
        <ellipse cx="100" cy="120" rx="65" ry="70" fill="url(#bodyGrad)" opacity="0.4"/>
        <defs>
            <linearGradient id="bodyGrad" x1="0" y1="0" x2="0" y2="1">
                <stop offset="0%" stop-color="#64ffda"/>
                <stop offset="100%" stop-color="#3a7bd5"/>
            </linearGradient>
        </defs>
        <circle cx="75" cy="105" r="22" fill="white"/>
        <circle cx="125" cy="105" r="22" fill="white"/>
        <circle cx="75" cy="105" r="10" fill="#1e1e2f"/>
        <circle cx="125" cy="105" r="10" fill="#1e1e2f"/>
        <circle cx="72" cy="101" r="3" fill="white"/>
        <circle cx="122" cy="101" r="3" fill="white"/>
        <polygon points="100,120 90,135 110,135" fill="#ffb74d"/>
        <path d="M 50 70 Q 100 30 150 70 Q 130 60 100 60 Q 70 60 50 70 Z" fill="#2c5364"/>
        <ellipse cx="60" cy="150" rx="15" ry="25" fill="#3a7bd5" transform="rotate(-20 60 150)"/>
        <ellipse cx="140" cy="150" rx="15" ry="25" fill="#3a7bd5" transform="rotate(20 140 150)"/>
        <rect x="70" y="155" width="60" height="40" rx="4" fill="#ffffff" opacity="0.9"/>
        <line x1="75" y1="165" x2="125" y2="165" stroke="#3a7bd5" stroke-width="2"/>
        <line x1="75" y1="175" x2="125" y2="175" stroke="#3a7bd5" stroke-width="2"/>
        <line x1="75" y1="185" x2="110" y2="185" stroke="#3a7bd5" stroke-width="2"/>
    </svg>
""", unsafe_allow_html=True)

mode = st.radio("Choose mode:", ["New Quiz", "Review Weak Questions", "Competitive Room"], horizontal=True)
import streamlit as st
import pdfplumber
import google.generativeai as genai
import json
import datetime

st.set_page_config(page_title="AI Study Buddy", page_icon="📚", layout="wide")

st.markdown("""""", unsafe_allow_html=True)

if mode == "Review Weak Questions":
    try:
        with open("quiz_history.json", "r") as f:
            history = json.load(f)
        all_wrong = []
        for attempt in history:
            all_wrong.extend(attempt['wrong_questions'])
        seen = set()
        review_questions = []
        for q in all_wrong:
            if q['question'] not in seen:
                review_questions.append(q)
                seen.add(q['question'])

        if not review_questions:
            st.success("🎉 No weak questions to review!")
        else:
            st.subheader(f"Reviewing {len(review_questions)} question(s)")
            review_score = 0
            for i, q in enumerate(review_questions):
                st.write(f"**Q{i+1}: {q['question']}**")
                options = [f"{k}) {v}" for k, v in q['options'].items()]
                choice = st.radio("Select:", options, key=f"rev{i}", index=None)
                if choice and choice[0] == q['correct_answer']:
                    review_score += 1
                st.divider()
            if st.button("Submit Review"):
                st.success(f"Review Score: {review_score}/{len(review_questions)}")
    except FileNotFoundError:
        st.info("No quiz history yet. Take a quiz first!")

    st.stop()

# ---------- COMPETITIVE ROOM MODE ----------
if mode == "Competitive Room":
    st.subheader("🏆 Competitive Room")
    
    room_action = st.radio("Do you want to:", ["Create a Room", "Join a Room"], horizontal=True)
    
    if room_action == "Create a Room":
        room_code = st.text_input("Set a Room Code (share this with friends):", value="ROOM123")
        your_name = st.text_input("Your Name:")
        uploaded_file = st.file_uploader("Upload notes to create quiz", type=["pdf", "docx", "txt"])
        num_q = st.number_input("Number of questions:", min_value=1, max_value=20, value=5)
        
        if uploaded_file and your_name and st.button("Create Room & Generate Quiz"):
            full_text = ""
            if uploaded_file.name.endswith(".pdf"):
                with pdfplumber.open(uploaded_file) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            full_text += text
            elif uploaded_file.name.endswith(".docx"):
                import docx
                doc = docx.Document(uploaded_file)
                for para in doc.paragraphs:
                    full_text += para.text + "\n"
            elif uploaded_file.name.endswith(".txt"):
                full_text = uploaded_file.read().decode("utf-8")

            genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
            model = genai.GenerativeModel("gemini-3.6-flash")

            prompt = f"""Generate {num_q} multiple choice quiz questions based on this text.
Return ONLY valid JSON, no other text, in this exact format:
[
  {{
    "question": "...",
    "options": {{"A": "...", "B": "...", "C": "...", "D": "..."}},
    "correct_answer": "A",
    "topic": "short topic name"
  }}
]
Text:
{full_text[:3000]}"""

            with st.spinner("Creating room..."):
                response = model.generate_content(prompt)
                clean_text = response.text.strip().replace("```json", "").replace("```", "")
                room_quiz = json.loads(clean_text)

                room_data = {
                    "quiz": room_quiz,
                    "leaderboard": []
                }
                with open(f"room_{room_code}.json", "w") as f:
                    json.dump(room_data, f, indent=2)

                st.success(f"✅ Room '{room_code}' created! Share this code with friends.")
                st.session_state.current_room = room_code
                st.session_state.room_quiz = room_quiz
                st.session_state.player_name = your_name

    elif room_action == "Join a Room":
        room_code = st.text_input("Enter Room Code:")
        your_name = st.text_input("Your Name:")
        
        if room_code and st.button("Join Room"):
            try:
                with open(f"room_{room_code}.json", "r") as f:
                    room_data = json.load(f)
                st.session_state.current_room = room_code
                st.session_state.room_quiz = room_data["quiz"]
                st.session_state.player_name = your_name
                st.success(f"Joined room '{room_code}'!")
            except FileNotFoundError:
                st.error("Room not found. Check the code and try again.")

    # If a room is active, show the quiz
    if "current_room" in st.session_state and "room_quiz" in st.session_state:
        st.divider()
        st.subheader(f"Room: {st.session_state.current_room}")
        quiz = st.session_state.room_quiz
        room_answers = {}

        for i, q in enumerate(quiz):
            st.write(f"**Q{i+1}: {q['question']}**")
            options = [f"{k}) {v}" for k, v in q['options'].items()]
            choice = st.radio("Select:", options, key=f"room_q{i}", index=None)
            if choice:
                room_answers[i] = choice[0]
            st.divider()

        if st.button("Submit & See Leaderboard"):
            score = sum(1 for i, q in enumerate(quiz) if room_answers.get(i) == q['correct_answer'])
            
            room_file = f"room_{st.session_state.current_room}.json"
            with open(room_file, "r") as f:
                room_data = json.load(f)
            
            room_data["leaderboard"].append({
                "name": st.session_state.player_name,
                "score": score,
                "total": len(quiz)
            })
            
            with open(room_file, "w") as f:
                json.dump(room_data, f, indent=2)

            st.success(f"Your Score: {score}/{len(quiz)}")

            st.subheader("🏆 Leaderboard")
            sorted_board = sorted(room_data["leaderboard"], key=lambda x: x["score"], reverse=True)
            for rank, entry in enumerate(sorted_board, 1):
                st.write(f"{rank}. **{entry['name']}** — {entry['score']}/{entry['total']}")

    st.stop()

# ---------- NEW QUIZ MODE ----------

uploaded_file = st.file_uploader("Upload your notes (PDF, Word, or Text)", type=["pdf", "docx", "txt"])
num_questions = st.number_input("How many questions?", min_value=1, max_value=100, value=5, step=1)

if uploaded_file is not None:
    if "quiz_data" not in st.session_state:
        full_text = ""

        if uploaded_file.name.endswith(".pdf"):
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text

        elif uploaded_file.name.endswith(".docx"):
            import docx
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                full_text += para.text + "\n"

        elif uploaded_file.name.endswith(".txt"):
            full_text = uploaded_file.read().decode("utf-8")

        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        model = genai.GenerativeModel("gemini-3.6-flash")

        prompt = f"""Generate {num_questions} multiple choice quiz questions based on this text.
Return ONLY valid JSON, no other text, in this exact format:
[
  {{
    "question": "...",
    "options": {{"A": "...", "B": "...", "C": "...", "D": "..."}},
    "correct_answer": "A",
    "topic": "short topic name"
  }}
]
Text:
{full_text[:3000]}"""

        with st.spinner("Generating quiz..."):
            response = model.generate_content(prompt)
            clean_text = response.text.strip().replace("```json", "").replace("```", "")
            st.session_state.quiz_data = json.loads(clean_text)
            st.session_state.answers = {}
            st.session_state.confidence = {}

    quiz_data = st.session_state.quiz_data

    for i, q in enumerate(quiz_data):
        st.subheader(f"Q{i+1}: {q['question']}")
        options = [f"{k}) {v}" for k, v in q['options'].items()]
        choice = st.radio("Select answer:", options, key=f"q{i}", index=None)
        confidence = st.select_slider(
            "How confident are you?",
            options=["Low", "Medium", "High"],
            key=f"conf{i}"
        )
        if choice:
            st.session_state.answers[i] = choice[0]
            st.session_state.confidence[i] = confidence
        st.divider()

    if st.button("Submit Quiz"):
        score = 0
        weak_topics = []
        overconfident = []
        underconfident = []
        confidences = st.session_state.get("confidence", {})

        for i, q in enumerate(quiz_data):
            user_ans = st.session_state.answers.get(i)
            conf = confidences.get(i, "Medium")
            is_correct = (user_ans == q['correct_answer'])

            if is_correct:
                score += 1
            else:
                weak_topics.append(q['topic'])

            if conf == "High" and not is_correct:
                overconfident.append(q['question'])
            if conf == "Low" and is_correct:
                underconfident.append(q['question'])

        st.success(f"Your Score: {score}/{len(quiz_data)}")

        if weak_topics:
            st.warning("Weak areas: " + ", ".join(set(weak_topics)))
        else:
            st.balloons()

        if overconfident:
            st.error("⚠️ Overconfident (sure, but wrong):")
            for q in overconfident:
                st.write(f"- {q}")

        if underconfident:
            st.info("💡 Underconfident (unsure, but correct):")
            for q in underconfident:
                st.write(f"- {q}")

        history_entry = {
            "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
            "score": score,
            "total": len(quiz_data),
            "weak_topics": weak_topics,
            "wrong_questions": [q for q in quiz_data if q['topic'] in weak_topics]
        }
        try:
            with open("quiz_history.json", "r") as f:
                history = json.load(f)
        except FileNotFoundError:
            history = []
        history.append(history_entry)
        with open("quiz_history.json", "w") as f:
            json.dump(history, f, indent=2)